#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INLINE_CODE_RE = re.compile(r"`([^`]+)`")


@dataclass(frozen=True)
class TableBlock:
    rows: list[list[str]]


def _set_document_fonts(document: Document) -> None:
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # East Asian font (Chinese). Word will fallback if the font is unavailable.
    rfonts = style.element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), "SimSun")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def _add_inline_runs(paragraph, text: str, *, monospace: str = "Consolas") -> None:
    if not text:
        return

    parts: list[tuple[str, bool]] = []
    last = 0
    for m in INLINE_CODE_RE.finditer(text):
        if m.start() > last:
            parts.append((text[last : m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))

    for content, is_code in parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        if is_code:
            run.font.name = monospace
            run._element.rPr.rFonts.set(qn("w:ascii"), monospace)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), monospace)


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|") or not s.endswith("|"):
        return False
    inner = s.strip("|").strip()
    if not inner:
        return False
    # Accept patterns like ---|---:|:---:
    return all(set(cell.strip()) <= set("-:") and "-" in cell for cell in inner.split("|"))


def _split_table_row(line: str) -> list[str]:
    s = line.strip()
    s = s.strip("|")
    return [cell.strip() for cell in s.split("|")]


def _parse_table_blocks(lines: list[str]) -> list[str | TableBlock]:
    out: list[str | TableBlock] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and line.strip().endswith("|"):
            # Need at least header + separator.
            if i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
                rows = [_split_table_row(line)]
                j = i + 2
                while j < len(lines) and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
                    rows.append(_split_table_row(lines[j]))
                    j += 1
                out.append(TableBlock(rows=rows))
                i = j
                continue
        out.append(line)
        i += 1
    return out


def _guess_numeric(s: str) -> bool:
    t = s.strip()
    if not t:
        return False
    # numeric-ish: 12,345 / 0.61% / 1.23 / +0.123 (+1.2%) / 4.05 MiB / 1.00x / abs≤11, rel≤0.006%
    if not any(ch.isdigit() for ch in t):
        return False
    cleaned = t
    cleaned = cleaned.replace(",", "")
    cleaned = cleaned.replace("MiB", "").replace("MB", "")
    cleaned = cleaned.replace("abs", "").replace("rel", "")
    cleaned = cleaned.replace("≤", "").replace("≥", "")
    cleaned = cleaned.replace("%", "").replace("×", "").replace("x", "")
    cleaned = cleaned.replace("/", "").replace("+", "").replace("-", "")
    cleaned = cleaned.replace("(", "").replace(")", "")
    cleaned = cleaned.replace(":", "").replace("=", "")
    cleaned = cleaned.replace(" ", "")
    cleaned = cleaned.replace(".", "")
    return cleaned.isdigit()


def _add_word_table(document: Document, block: TableBlock) -> None:
    if not block.rows:
        return
    ncols = max(len(r) for r in block.rows)
    table = document.add_table(rows=len(block.rows), cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    header = [block.rows[0][i] if i < len(block.rows[0]) else "" for i in range(ncols)]
    header_lc = [h.strip().lower() for h in header]
    prefer_right = set()
    for i, h in enumerate(header_lc):
        if any(
            k in h
            for k in [
                "paper",
                "ours",
                "Δ".lower(),
                "models",
                "model",
                "space",
                "mib",
                "memory",
                "rp",
                "reduction",
                "ratio",
                "avg",
                "%",  # reduction %
            ]
        ):
            prefer_right.add(i)

    for r_i, row in enumerate(block.rows):
        for c_i in range(ncols):
            text = row[c_i] if c_i < len(row) else ""
            cell = table.cell(r_i, c_i)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, text)

            if r_i == 0:
                for run in p.runs:
                    run.bold = True
            else:
                if c_i in prefer_right or _guess_numeric(text):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a little spacing after tables
    document.add_paragraph("")


def _add_code_block(document: Document, code_lines: Iterable[str]) -> None:
    p = document.add_paragraph()
    run = p.add_run("".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    p.paragraph_format.space_after = Pt(6)


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    raw_lines = md_path.read_text(encoding="utf-8").splitlines(keepends=False)
    blocks = _parse_table_blocks(raw_lines)

    doc = Document()
    _set_document_fonts(doc)

    in_code = False
    code_buf: list[str] = []

    for item in blocks:
        if isinstance(item, TableBlock):
            if in_code:
                _add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            _add_word_table(doc, item)
            continue

        line = item.rstrip("\n")

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                _add_code_block(doc, code_buf)
                code_buf = []
            continue

        if in_code:
            code_buf.append(line + "\n")
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            level = max(1, min(level, 4))
            p = doc.add_paragraph(style=f"Heading {level}")
            _add_inline_runs(p, title)
            continue

        # Blockquote
        if line.lstrip().startswith(">"):
            content = line.lstrip()[1:].lstrip()
            p = doc.add_paragraph(style="Intense Quote" if "Intense Quote" in doc.styles else None)
            if p.style is None:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
            _add_inline_runs(p, content)
            continue

        # Bullet / ordered lists (simple)
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if m:
            indent, marker, content = m.groups()
            level = 2 if len(indent) >= 2 else 1
            if marker.endswith(".") and marker[:-1].isdigit():
                style = "List Number 2" if level == 2 and "List Number 2" in doc.styles else "List Number"
            else:
                style = "List Bullet 2" if level == 2 and "List Bullet 2" in doc.styles else "List Bullet"
            p = doc.add_paragraph(style=style)
            _add_inline_runs(p, content)
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline_runs(p, line)

    if in_code and code_buf:
        _add_code_block(doc, code_buf)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    ap = argparse.ArgumentParser(description="Convert a Markdown report to DOCX with real Word tables.")
    ap.add_argument("md", type=Path, help="Input Markdown file path.")
    ap.add_argument("out", type=Path, help="Output DOCX file path.")
    args = ap.parse_args()

    convert_md_to_docx(args.md, args.out)


if __name__ == "__main__":
    main()
